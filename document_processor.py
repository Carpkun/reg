import os
import re
from docx import Document
from typing import List, Dict, Tuple
from langchain.text_splitter import RecursiveCharacterTextSplitter

class DocumentProcessor:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 150):
        """
        문서 처리기 초기화
        
        Args:
            chunk_size: 텍스트 청크 크기
            chunk_overlap: 청크 간 중복 길이
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        docx 파일에서 텍스트 추출
        
        Args:
            file_path: docx 파일 경로
            
        Returns:
            추출된 텍스트
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # 문서 제목 추가
            if doc.core_properties.title:
                text_content.append(f"제목: {doc.core_properties.title}")
            
            # 문서 본문 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 표 내용 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            return self.clean_text(full_text)
            
        except Exception as e:
            print(f"문서 읽기 오류 ({file_path}): {str(e)}")
            return ""
    
    def clean_text(self, text: str) -> str:
        """
        텍스트 정제
        
        Args:
            text: 원본 텍스트
            
        Returns:
            정제된 텍스트
        """
        # 연속된 공백과 줄바꿈 정리
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r' +', ' ', text)
        
        # 앞뒤 공백 제거
        text = text.strip()
        
        return text
    
    def process_documents(self, data_folder: str) -> List[Dict[str, str]]:
        """
        data 폴더의 모든 docx 파일 처리
        
        Args:
            data_folder: 문서가 저장된 폴더 경로
            
        Returns:
            처리된 문서 청크 리스트
        """
        documents = []
        
        print(f"🔍 [DEBUG] 문서 처리 시작: {data_folder}")
        print(f"🔍 [DEBUG] 절대 경로: {os.path.abspath(data_folder)}")
        
        if not os.path.exists(data_folder):
            print(f"❌ [ERROR] {data_folder} 폴더가 존재하지 않습니다.")
            return documents
        
        try:
            all_files = os.listdir(data_folder)
            print(f"🔍 [DEBUG] 폴더 내 전체 파일 수: {len(all_files)}")
            
            docx_files = [f for f in all_files if f.endswith('.docx')]
            print(f"🔍 [DEBUG] docx 파일 수: {len(docx_files)}")
            
            if docx_files:
                print(f"🔍 [DEBUG] 첫 번째 docx 파일: {docx_files[0]}")
                print(f"🔍 [DEBUG] 처음 5개 파일: {docx_files[:5]}")
            
        except Exception as e:
            print(f"❌ [ERROR] 폴더 읽기 오류: {str(e)}")
            return documents
        
        if not docx_files:
            print(f"⚠️ [WARNING] {data_folder} 폴더에 docx 파일이 없습니다.")
            return documents
        
        print(f"📚 [INFO] 총 {len(docx_files)}개의 docx 파일을 처리합니다...")
        
        for i, filename in enumerate(docx_files):
            file_path = os.path.join(data_folder, filename)
            print(f"📄 [INFO] 처리 중 ({i+1}/{len(docx_files)}): {filename}")
            
            try:
                # 파일 존재 및 읽기 권한 확인
                if not os.path.exists(file_path):
                    print(f"❌ [ERROR] 파일이 존재하지 않음: {file_path}")
                    continue
                
                if not os.access(file_path, os.R_OK):
                    print(f"❌ [ERROR] 파일 읽기 권한 없음: {file_path}")
                    continue
                
                # 파일 크기 확인
                file_size = os.path.getsize(file_path)
                print(f"🔍 [DEBUG] 파일 크기: {file_size} bytes")
                
                # 텍스트 추출
                content = self.extract_text_from_docx(file_path)
                
                if content:
                    print(f"✅ [SUCCESS] 텍스트 추출 성공, 길이: {len(content)} 문자")
                    
                    # 텍스트를 청크로 분할
                    chunks = self.text_splitter.split_text(content)
                    
                    for j, chunk in enumerate(chunks):
                        documents.append({
                            'content': chunk,
                            'metadata': {
                                'source': filename,
                                'chunk_id': j,
                                'file_path': file_path
                            }
                        })
                    
                    print(f"✅ [SUCCESS] {len(chunks)}개 청크 생성")
                else:
                    print(f"❌ [ERROR] 텍스트 추출 실패: {filename}")
                    
            except Exception as e:
                print(f"❌ [ERROR] 파일 처리 오류 ({filename}): {str(e)}")
                print(f"🔍 [DEBUG] 오류 타입: {type(e).__name__}")
                import traceback
                print(f"🔍 [DEBUG] 스택 트레이스: {traceback.format_exc()}")
        
        print(f"🎉 [FINAL] 총 {len(documents)}개의 문서 청크가 생성되었습니다.")
        return documents
    
    def get_document_stats(self, documents: List[Dict[str, str]]) -> Dict[str, int]:
        """
        문서 통계 정보 반환
        
        Args:
            documents: 처리된 문서 리스트
            
        Returns:
            통계 정보 딕셔너리
        """
        if not documents:
            return {'총_문서수': 0, '총_청크수': 0, '평균_청크길이': 0}
        
        total_chunks = len(documents)
        total_length = sum(len(doc['content']) for doc in documents)
        avg_chunk_length = total_length // total_chunks if total_chunks > 0 else 0
        
        # 고유한 소스 파일 개수
        unique_sources = len(set(doc['metadata']['source'] for doc in documents))
        
        return {
            '총_문서수': unique_sources,
            '총_청크수': total_chunks,
            '평균_청크길이': avg_chunk_length,
            '총_문자수': total_length
        } 